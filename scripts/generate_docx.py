#!/usr/bin/env python3
"""Generate Oligool_Features.docx from the report content, matching the PDF style.

PDF style reference (via pdffonts): Arial + ArialBold, A4 page size.
Heading sizes match pandoc defaults: H1=20pt, H2=16pt, body=12pt.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Arial"
BODY_SIZE = Pt(12)
H1_SIZE = Pt(20)
H2_SIZE = Pt(16)
TITLE_SIZE = Pt(28)
BULLET_INDENT = Cm(1.0)
SUB_INDENT = Cm(2.0)


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, monospace=False):
    """Apply font settings to a run, defaulting to Arial."""
    font_name = "Courier New" if monospace else FONT_NAME
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    # Ensure the font applies for East Asian / complex script too
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def add_paragraph(doc, text="", size=BODY_SIZE, bold=False, alignment=None,
                  space_before=Pt(0), space_after=Pt(6)):
    """Add a simple paragraph with uniform formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def add_horizontal_rule(doc):
    """Add a horizontal rule as a bottom-border on an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")        # 0.75pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def add_numbered_item(doc, number, segments, sub_indent=BULLET_INDENT):
    """Add a numbered list item.

    Args:
        number: The list number (int).
        segments: List of (text, bold, monospace) tuples for the lead-in line.
        sub_indent: Indentation for the numbered item.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = sub_indent
    p.paragraph_format.first_line_indent = Cm(-0.6)  # hanging indent for number
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

    num_run = p.add_run(f"{number}. ")
    set_run_font(num_run, size=BODY_SIZE, bold=True)

    for text, bold, monospace in segments:
        run = p.add_run(text)
        set_run_font(run, size=BODY_SIZE, bold=bold, monospace=monospace)
    return p


def add_sub_bullet(doc, segments, indent=SUB_INDENT):
    """Add a sub-bullet (dash) under a numbered item.

    Args:
        segments: List of (text, bold, monospace) tuples.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Cm(-0.5)  # hanging indent for dash
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

    dash_run = p.add_run("\u2013  ")  # en-dash
    set_run_font(dash_run, size=BODY_SIZE)

    for text, bold, monospace in segments:
        run = p.add_run(text)
        set_run_font(run, size=BODY_SIZE, bold=bold, monospace=monospace)
    return p


def add_plain_sub(doc, text, indent=SUB_INDENT):
    """Add a plain sub-bullet (no bold lead-in, just dash + text)."""
    return add_sub_bullet(doc, [(text, False, False)], indent=indent)


def build_document():
    doc = Document()

    # --- Page setup: A4 ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)     # A4 width
    section.page_height = Cm(29.7)    # A4 height
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Set Normal style to Arial ---
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)

    # --- Title block (from YAML metadata) ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(0)
    title_run = title_p.add_run("Oligool: Project Features & Workflow")
    set_run_font(title_run, size=TITLE_SIZE, bold=True)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(0)
    author_run = author_p.add_run("Mgr. Vojtěch Rejtar")
    set_run_font(author_run, size=Pt(14))

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(12)
    date_run = date_p.add_run("February 24, 2026")
    set_run_font(date_run, size=Pt(14))

    # --- H1: Oligool: Project Overview ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1_run = h1.add_run("Oligool: Project Overview")
    set_run_font(h1_run, size=H1_SIZE, bold=True)

    # --- Author / Date / Project block ---
    info_p = doc.add_paragraph()
    info_p.paragraph_format.space_after = Pt(6)
    for label, value in [("Author", "Mgr. Vojtěch Rejtar"),
                         ("Date", "February 24, 2026"),
                         ("Project", "Oligool")]:
        r1 = info_p.add_run(f"{label}: ")
        set_run_font(r1, size=BODY_SIZE, bold=True)
        r2 = info_p.add_run(value)
        set_run_font(r2, size=BODY_SIZE)
        if label != "Project":
            sep = info_p.add_run("  ")
            set_run_font(sep, size=BODY_SIZE)

    # --- Intro paragraph ---
    add_paragraph(
        doc,
        "Oligool is a native desktop application designed to streamline the "
        "design, alignment, and analysis of genetic sequences and custom oligos "
        "for molecular biologists.",
        space_before=Pt(6),
        space_after=Pt(6),
    )

    # --- Horizontal rule ---
    add_horizontal_rule(doc)

    # --- H2: Features and Capabilities ---
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("Features and Capabilities")
    set_run_font(h2_run, size=H2_SIZE, bold=True)

    # 1. Dual-Source Search & Fetch
    add_numbered_item(doc, 1, [("Dual-Source Search & Fetch", True, False)])
    add_plain_sub(doc, "Seamless toggling between NCBI and Ensembl data sources "
                       "for finding genes and transcripts.")
    add_plain_sub(doc, "Smart filtering with persistent search parameters "
                       "(E-value, Identity %, Organism) that survive application restarts.")

    # 2. Interactive MSA Viewer
    add_numbered_item(doc, 2, [("Interactive MSA Viewer", True, False)])
    add_plain_sub(doc, "High-performance Multiple Sequence Alignment powered by MAFFT.")
    add_plain_sub(doc, "2D Navigation via an interactive minimap for scrubbing through "
                       "massive alignments, highlighting conservation patterns, and "
                       "identifying variations.")

    # 3. "Oligize!" Design
    add_numbered_item(doc, 3, [("\u201cOligize!\u201d Design", True, False)])
    add_plain_sub(doc, "Precision splitting of genomic regions into two contiguous "
                       "oligos with exact control over shift and lengths.")
    add_plain_sub(doc, "Live integration with IDT OligoAnalyzer to evaluate hairpin "
                       "formation and self-dimerization in real time (Delta G).")

    # 4. "Primerize!" Schematic
    add_numbered_item(doc, 4, [("\u201cPrimerize!\u201d Schematic", True, False)])
    add_plain_sub(doc, "High-fidelity SVG visual assembly of your molecular design, "
                       "accurately representing Forward and Reverse Primer Binding "
                       "Sites (PBS) and TAG sequences.")
    add_sub_bullet(doc, [
        ("Seq Mode: ", True, False),
        ("Toggle to a high-detail view displaying base-by-base lettering along "
         "the schematics architecture.", False, False),
    ])
    add_sub_bullet(doc, [
        ("Persistence: ", True, False),
        ("Local storage of user TAGs, PBS sequences, and design preferences.", False, False),
    ])

    # --- Horizontal rule ---
    add_horizontal_rule(doc)

    # --- H2: Typical Workflow ---
    h2b = doc.add_paragraph()
    h2b.paragraph_format.space_before = Pt(6)
    h2b.paragraph_format.space_after = Pt(6)
    h2b_run = h2b.add_run("Typical Workflow")
    set_run_font(h2b_run, size=H2_SIZE, bold=True)

    # 1. Launch the Application
    add_numbered_item(doc, 1, [("Launch the Application", True, False)])
    add_sub_bullet(doc, [
        ("macOS: ", True, False),
        ("Run the standalone Mac bundle ", False, False),
        ("Oligool.app", False, True),
        (".", False, False),
    ])
    add_sub_bullet(doc, [
        ("Windows: ", True, False),
        ("Execute the single-file executable ", False, False),
        ("Oligool.exe", False, True),
        (".", False, False),
    ])
    add_plain_sub(doc, "All credentials (e.g., NCBI Key, IDT API authentication) and "
                       "design configurations are securely stored locally via ")
    # The last sub-bullet needs localStorage as monospace — redo as segments
    # Remove the plain sub and re-add with monospace tail
    doc.paragraphs[-1].clear() if hasattr(doc.paragraphs[-1], 'clear') else None
    # Actually python-docx doesn't have .clear(); remove last paragraph element
    last_p = doc.paragraphs[-1]
    last_p._element.getparent().remove(last_p._element)
    add_sub_bullet(doc, [
        ("All credentials (e.g., NCBI Key, IDT API authentication) and design "
         "configurations are securely stored locally via ", False, False),
        ("localStorage", False, True),
        (".", False, False),
    ])

    # 2. Search and Retrieve Sequences
    add_numbered_item(doc, 2, [("Search and Retrieve Sequences", True, False)])
    add_plain_sub(doc, "Toggle between NCBI and Ensembl.")
    add_plain_sub(doc, "Enter your target gene name or accession ID to retrieve "
                       "annotated sequence data.")

    # 3. Sequence Alignment (Optional)
    add_numbered_item(doc, 3, [("Sequence Alignment (Optional)", True, False)])
    add_plain_sub(doc, "Upload multiple sequences and utilize the MAFFT-powered MSA "
                       "Viewer to align them.")
    add_plain_sub(doc, "Use the 2D minimap to identify conserved regions.")

    # 4. Oligo Design ("Oligize!")
    add_numbered_item(doc, 4, [("Oligo Design (\u201cOligize!\u201d)", True, False)])
    add_plain_sub(doc, "Select a genomic region of interest.")
    add_plain_sub(doc, "Define your desired shift and oligo lengths to split the "
                       "region into two contiguous oligos.")
    add_plain_sub(doc, "Review the live IDT OligoAnalyzer real-time Delta G values "
                       "to avoid hairpins and self-dimerization.")

    # 5. Visual Assembly ("Primerize!")
    add_numbered_item(doc, 5, [("Visual Assembly (\u201cPrimerize!\u201d)", True, False)])
    add_plain_sub(doc, "Navigate to the schematic view to visually assemble your "
                       "molecular design.")
    add_plain_sub(doc, "Input your custom TAGs, Forward PBS, and Reverse PBS.")
    add_sub_bullet(doc, [
        ("Enable ", False, False),
        ("Seq Mode", True, False),
        (" to verify the base-by-base construct architecture.", False, False),
    ])
    add_plain_sub(doc, "Use the generated sequence constructs for your downstream "
                       "experimental workflows.")

    # --- Set heading styles font to Arial as well ---
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            rpr = style.element.get_or_add_rPr()
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rpr.append(rf)
            rf.set(qn("w:ascii"), FONT_NAME)
            rf.set(qn("w:hAnsi"), FONT_NAME)
            rf.set(qn("w:cs"), FONT_NAME)
        except KeyError:
            pass

    return doc


if __name__ == "__main__":
    doc = build_document()
    output = "Oligool_Features.docx"
    doc.save(output)
    print(f"Generated {output}")
